from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "AI-Test-Officer-使用说明.docx"
ASSETS = ROOT / "assets"

GREEN = RGBColor(24, 53, 45)
INK = RGBColor(27, 35, 31)
MUTED = RGBColor(92, 106, 97)
SOFT = "EAF1EB"
LINE = "D8E1DA"


def set_font(run, name="Aptos", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def shade(paragraph, fill):
    props = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_cell_padding(cell, top=110, start=160, bottom=110, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_text(doc, text, style="Normal", before=None, after=None, color=None, size=None, bold=None):
    p = doc.add_paragraph(style=style)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    shade(p, SOFT)
    label_run = p.add_run(f"{label}  ")
    set_font(label_run, size=10.5, color=GREEN, bold=True)
    body = p.add_run(text)
    set_font(body, size=10.5, color=INK)
    return p


def add_figure(doc, filename, caption, width=6.45):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(ASSETS / filename), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(9)
    r = caption_p.add_run(caption)
    set_font(r, size=9, color=MUTED)
    return p


def add_step(doc, number, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{number}. {title}")
    set_font(r, size=14, color=GREEN, bold=True)
    body_p = doc.add_paragraph()
    body_p.paragraph_format.space_after = Pt(6)
    body_p.paragraph_format.left_indent = Inches(0.1)
    body_r = body_p.add_run(body)
    set_font(body_r, size=10.5, color=INK)


def add_module(doc, title, purpose, action):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, size=11, color=GREEN, bold=True)
    add_text(doc, f"作用：{purpose}", before=0, after=1, size=10.3, color=INK)
    add_text(doc, f"你要做：{action}", before=0, after=5, size=10.3, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, before, after in (("Heading 1", 18, 16, 6), ("Heading 2", 13, 12, 4), ("Heading 3", 11, 8, 3)):
        style = styles[name]
        style.font.name = "Aptos Display" if name == "Heading 1" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.color.rgb = GREEN
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("AI Test Officer  |  新手使用说明")
    set_font(header_run, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("本地测试工作台 · 证据、结论与报告均在同一次运行中关联")
    set_font(footer_run, size=8, color=MUTED)


def build():
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("AI TEST OFFICER")
    set_font(r, size=10, color=GREEN, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("零基础使用说明")
    set_font(r, name="Aptos Display", size=30, color=GREEN, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("从接入一个本地项目，到拿到一份可追溯的测试结论。")
    set_font(r, size=13, color=MUTED)
    add_callout(doc, "先记住", "这不是让你手写测试脚本的工具。你告诉它要测哪个项目、参考哪些需求或改动，它会生成计划、操作浏览器、收集证据，再告诉你能不能继续发布。")
    add_text(doc, "你只需要完成这 4 件事", style="Heading 2")
    for title, text in (
        ("接入项目", "选择本地项目文件夹，让系统识别技术栈、启动命令和端口。"),
        ("给出依据", "添加需求、PR、Bug 单或 OpenAPI；系统会说明本次测试参考了什么。"),
        ("确认计划", "选择现成场景，或让 AI 发现测试点并生成待审核草案。"),
        ("运行并处理结论", "授权浏览器后开始；查看截图、日志、断言、失败原因和下一步建议。"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{title}：")
        set_font(r1, size=10.5, color=GREEN, bold=True)
        r2 = p.add_run(text)
        set_font(r2, size=10.5, color=INK)

    doc.add_page_break()
    add_text(doc, "1. 先认识主界面", style="Heading 1")
    add_text(doc, "打开工作台后，不需要理解所有技术词。先把画面当作一个测试任务驾驶舱：左边决定测什么，中间显示 AI 正在做什么，右边告诉你结果和下一步。", after=8)
    add_figure(doc, "01-workbench-main.png", "图 1：新版工作台首屏。测试核心信息集中在一个任务舞台中。")
    add_module(doc, "左侧：本次测试", "按顺序进入项目、测试依据和影响分析。这里是准备区，不是日志区。", "第一次使用时，先点 01 项目与连接配置；项目能连通后，再点 02 读取来源、03 分析影响。")
    add_module(doc, "中间：AI 测试任务", "最重要的区域。它显示为什么测这个功能、测试对象、依据数量、计划步数和 Agent 时间线。", "确认任务标题和“为什么测”符合你的意图；勾选浏览器授权后点击“开始测试”。")
    add_module(doc, "右侧：测试结论", "显示是否通过、收集到多少证据、下一步建议，以及提交检查和需求验收入口。", "运行结束先看结论，再点“查看完整证据”了解截图、请求、日志和断言。")

    doc.add_page_break()
    add_text(doc, "2. 接入要测试的项目", style="Heading 1")
    add_text(doc, "这是第一次使用时最需要完成的设置。点击顶部“详细配置”，左侧会打开项目接入向导。它不会把密码写进项目配置；账号和密钥使用环境变量或凭据引用。", after=8)
    add_figure(doc, "02-project-onboarding-inputs.png", "图 2：项目接入向导。左侧抽屉只在配置时打开，避免打断日常测试。")
    add_step(doc, "第一步", "填写项目文件夹", "把要测试的 React、Vite、Next、FastAPI 或 Express 项目所在文件夹填入“项目文件夹”。若项目在当前测试平台目录之外，勾选“允许外部绝对路径”。")
    add_step(doc, "第二步", "点击自动识别", "系统会猜测项目技术栈、常用安装与启动命令、前端/后端端口和健康检查地址。你可以点击“套用建议”，再按项目实际情况修改。")
    add_step(doc, "第三步", "测试连接", "点击“测试连接”。成功表示系统能访问项目；失败会告诉你问题属于安装、启动、端口、健康检查或缺少登录账号中的哪一种，并给出白话修复提示。")
    add_callout(doc, "小白判断标准", "只要“测试连接”通过，并且浏览器能打开前端地址，就说明这个项目已经可以成为测试对象。无需先理解 Playwright 或 API。")

    # Let the next section start on the following natural page boundary.
    add_text(doc, "3. 告诉系统：这次为什么要测", style="Heading 1")
    add_text(doc, "测试依据决定 AI 应该关注哪里。你可以用一个或多个来源，不需要全部都有。真实来源读取失败时，系统会明确显示 missing 或权限不足，不会偷偷把演示数据当成你的真实输入。", after=7)
    add_module(doc, "需求 Markdown / 本地文件", "让系统知道产品期望的行为，例如“提交后必须显示成功提示”。", "在详细配置的“需求”或 MCP Connector 输入中填文件路径或 URL。")
    add_module(doc, "GitHub PR / Git diff", "让系统知道代码最近改了什么，从而推荐受影响的页面和测试场景。", "填 PR 地址、Diff 地址或粘贴本地 diff；读取后点“分析输入”。")
    add_module(doc, "Issue、TAPD、Bug URL", "让系统带着已知问题做回归，避免修复后再次出现。", "填写 Bug 单路径或 URL；若没有访问权限，报告会直接说明。")
    add_module(doc, "OpenAPI 文档", "让系统了解接口路径、参数与返回结构，可用于接口失败和契约验证。", "填 OpenAPI JSON/YAML 文件路径或 URL。")
    add_callout(doc, "建议", "接真实项目时勾选 strictInput。这样远程来源读不到时，测试会停下来提示你补充权限，而不会使用本地演示 fixture。")
    add_text(doc, "读取后会发生什么", style="Heading 2")
    add_text(doc, "点击“读取来源”后，系统记录每个来源的类型、读取状态、权限状态、内容摘要、可信等级、读取时间和内容哈希。点击“分析影响”后，系统将改动文件、路由、接口和关键词映射为推荐测试场景；没有覆盖到的内容会进入 harness gap，明确告诉你目前缺少什么测试能力。")

    doc.add_page_break()
    add_text(doc, "4. 选择测试场景，或让 AI 提出新测试点", style="Heading 1")
    add_text(doc, "左下角“测试场景”是可直接执行的能力库，例如登录、权限拦截、列表筛选、创建与编辑、表单校验、接口失败恢复、文件上传、OpenAPI 契约等。选择一个场景后，中间会显示本次要执行的计划。", after=7)
    add_step(doc, "已有场景", "直接选择并运行", "适合你已经知道要验证什么。例如修改了登录接口，就选择登录成功、登录失败或权限拦截。")
    add_step(doc, "AI 新建议", "先生成草案再审核", "在“更多能力”中使用 Discovery。它会先读取页面按钮、表单、data-testid、可访问链接和接口，再提出测试点。新的测试点不会自动执行或自动写进能力库。")
    add_step(doc, "草案入库", "探测后人工批准", "系统先验证选择器能否找到元素、断言是否有证据依据。探测和 oracle dry-run 都通过后，你批准草案，它才会成为以后可复用的场景。")
    add_callout(doc, "为什么要多一步批准", "这样可以防止 AI 因为猜错按钮、猜错页面或证据不足而执行不可靠的测试。平台宁可把它标成“需要补充能力”，也不会假装已经覆盖。")

    doc.add_page_break()
    add_text(doc, "5. 开始测试后，你该看什么", style="Heading 1")
    add_text(doc, "勾选“允许本次会话接管指定浏览器窗口执行测试”，再点击绿色“开始测试”。测试不会只给一个成功或失败；每一步都会对应计划、动作、断言和证据。", after=8)
    add_text(doc, "中间时间线", style="Heading 2")
    add_text(doc, "这里按顺序显示 Agent 做的事：生成计划、打开页面、点击或输入、等待接口、执行断言、重试、生成报告。它回答的是“AI 现在在做什么”。")
    add_text(doc, "右侧测试结论", style="Heading 2")
    add_text(doc, "这里先给白话结果：通过、失败、阻塞、可能不稳定或证据不足。下方“下一步建议”会告诉你优先检查项目连接、需求、代码、接口还是测试脚本。")
    add_text(doc, "证据详情", style="Heading 2")
    add_text(doc, "点击顶部“证据详情”或右侧“查看完整证据”，可以打开完整审计面板。你会看到截图、DOM 快照、console 报错、network 请求、trace/video、断言结果、Judge 结论和失败归因。每条 Judge 结论都要引用真实 evidence ID；没有证据时只能输出“证据不足”。")
    add_callout(doc, "快速定位失败", "先看失败步骤，再看对应截图和网络请求。若结论提供 top suspect，就继续看它关联的 diff 文件、组件名、接口地址与证据编号；这比只看“测试失败”更容易决定该找谁处理。")

    doc.add_page_break()
    add_text(doc, "6. 运行结束后：怎么处理和复用", style="Heading 1")
    add_module(doc, "通过", "说明当前测试计划中的断言都被证据支持。", "点击提交检查或需求验收，生成可被 CI 消费的 gate/report 产物。")
    add_module(doc, "失败", "可能是产品问题、测试脚本失效、环境问题或证据不足。", "打开证据详情，按“失败步骤 -> 截图/日志/网络 -> 下一步建议”的顺序排查。")
    add_module(doc, "阻塞", "通常是项目没启动、端口不通、账号缺失、权限不足或真实来源无法读取。", "回到详细配置的项目向导或 Connector 输入，按提示补齐后重新运行。")
    add_module(doc, "长期巡检", "可把核心场景设为定时巡检，追踪历史趋势、重试和风险升高提醒。", "在“更多能力”配置 Patrol；有失败时可通过企业微信、飞书、Slack 或 GitHub PR 评论推送。")
    add_text(doc, "安全边界", style="Heading 2")
    add_text(doc, "不要把账号密码、cookie、token 或 webhook 链接直接写进需求、报告和截图说明中。项目配置只保存环境变量或 Credential ID 引用；平台会在 UI、日志、报告与机器人消息中脱敏敏感字段。生产或真实接入环境需要显式 token，不能依赖 dev-local-token。")
    add_text(doc, "最短使用路径", style="Heading 2")
    add_text(doc, "配置项目 -> 测试连接 -> 读取来源 -> 分析影响 -> 选择场景 -> 勾选浏览器授权 -> 开始测试 -> 查看结论与完整证据。第一次使用只要按这个顺序走，就能完成一次完整测试。")

    doc.save(OUT)


if __name__ == "__main__":
    build()
